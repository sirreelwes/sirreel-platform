#!/usr/bin/env python3
"""
Generates a scaffold .docx at public/contracts/sirreel-rental-agreement-template.docx
with all 16 placeholders the portal Path A download endpoint expects. This
scaffold is NOT the canonical rental agreement — it is a structural placeholder
so the docxtemplater pipeline can be tested end-to-end. Before this feature
ships to clients, the body of this .docx must be replaced with the canonical
agreement content (post the corrections in docs/specs/canonical-baseline-
corrections.md), preserving the {{placeholder}} tokens in the same positions
they appear in the canonical PDF.

Regenerate with:
    python3 scripts/generate-agreement-template.py
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

OUT_PATH = Path(__file__).resolve().parent.parent / "public" / "contracts" / "sirreel-rental-agreement-template.docx"


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Helvetica"


def add_paragraph(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Helvetica"
    run.font.size = Pt(11)
    run.bold = bold


def add_field(doc: Document, label: str, placeholder: str) -> None:
    p = doc.add_paragraph()
    label_run = p.add_run(f"{label}: ")
    label_run.bold = True
    label_run.font.name = "Helvetica"
    label_run.font.size = Pt(11)
    value_run = p.add_run(placeholder)
    value_run.font.name = "Helvetica"
    value_run.font.size = Pt(11)


def main() -> None:
    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # ── Banner ──────────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("SIRREEL STUDIO RENTALS")
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.name = "Helvetica"

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run("Equipment and Vehicle Rental Agreement")
    sub_run.italic = True
    sub_run.font.size = Pt(12)
    sub_run.font.name = "Helvetica"

    addr = doc.add_paragraph()
    addr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    addr_run = addr.add_run("8500 Lankershim Blvd, Sun Valley, CA 91352 · (888) 477-7335")
    addr_run.font.size = Pt(10)
    addr_run.font.name = "Helvetica"

    doc.add_paragraph()

    # ── Template scaffold notice ────────────────────────────────────────
    note = doc.add_paragraph()
    note_run = note.add_run(
        "TEMPLATE SCAFFOLD — Replace this body with the canonical rental "
        "agreement content (post-corrections per canonical-baseline-"
        "corrections.md) before shipping to clients. Preserve every "
        "{{placeholder}} token in its canonical position so docxtemplater can "
        "fill them at runtime."
    )
    note_run.italic = True
    note_run.font.size = Pt(9)
    note_run.font.name = "Helvetica"

    doc.add_paragraph()

    # ── Job header ──────────────────────────────────────────────────────
    add_heading(doc, "Agreement Details", level=2)
    add_field(doc, "Generated", "{{generatedDate}}")
    add_field(doc, "Job", "{{jobName}} ({{jobNumber}})")
    add_field(doc, "Job type", "{{jobType}}")
    add_field(doc, "Rental period", "{{rentalStart}} – {{rentalEnd}}")

    doc.add_paragraph()

    # ── Lessee block ────────────────────────────────────────────────────
    add_heading(doc, "Lessee / Renter", level=2)
    add_field(doc, "Company name", "{{companyName}}")
    add_field(doc, "Company type", "{{companyType}}")
    add_field(doc, "Address", "{{companyAddress}}")
    add_field(doc, "Email", "{{companyEmail}}")
    add_field(doc, "Phone", "{{companyPhone}}")

    doc.add_paragraph()

    # ── Production contact ──────────────────────────────────────────────
    add_heading(doc, "Production Contact", level=2)
    add_field(doc, "Name", "{{contactFirstName}} {{contactLastName}}")
    add_field(doc, "Title", "{{contactPosition}}")
    add_field(doc, "Email", "{{contactEmail}}")
    add_field(doc, "Phone", "{{contactPhone}}")

    doc.add_paragraph()

    # ── Body placeholder ────────────────────────────────────────────────
    add_heading(doc, "Terms and Conditions", level=2)
    add_paragraph(
        doc,
        "Replace this paragraph with the canonical numbered clauses (1–29) "
        "plus the Fleet Agreement and LCDW addendum from the post-corrections "
        "rental agreement. Keep this section's placeholders ({{companyName}}, "
        "{{contactEmail}}, etc.) where the canonical PDF uses them so a filled "
        ".docx mirrors the legal PDF exactly.",
    )

    doc.add_paragraph()

    # ── Signature block (Path A clients sign offline; portal Path B signs natively) ──
    add_heading(doc, "Signature", level=2)
    add_paragraph(doc, "Authorized signatory:", bold=True)
    add_field(doc, "Name", "{{contactFirstName}} {{contactLastName}}")
    add_field(doc, "Title", "{{contactPosition}}")
    add_paragraph(doc, "Signature: ____________________________   Date: ____________")

    doc.save(OUT_PATH)
    print(f"wrote {OUT_PATH}")


if __name__ == "__main__":
    main()
